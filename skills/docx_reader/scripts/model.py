from docx import Document
import os
from typing import Dict, List, Optional

class DOCXReader:
    def __init__(self, file_path: str):
        """
        初始化DOCX读取器
        :param file_path: DOCX文件的绝对路径
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        if not file_path.lower().endswith('.docx'):
            raise ValueError("输入文件必须是DOCX格式")
        
        self.file_path = file_path
        self.document = None
        self._open_docx()

    def _open_docx(self):
        """
        内部方法：打开DOCX文件
        """
        try:
            self.document = Document(self.file_path)
        except Exception as e:
            raise RuntimeError(f"无法打开DOCX文件: {str(e)}")

    def extract_text(self) -> str:
        """
        提取DOCX文件的全部文本内容
        :return: 包含所有文本的字符串
        """
        if not self.document:
            self._open_docx()
        
        text = ""
        for paragraph in self.document.paragraphs:
            text += paragraph.text + "\n"
        
        return text.strip()

    def extract_metadata(self) -> Dict:
        """
        提取DOCX文件的元数据信息
        :return: 包含元数据的字典
        """
        if not self.document:
            self._open_docx()
        
        core_properties = self.document.core_properties
        metadata = {
            'title': core_properties.title or '',
            'author': core_properties.author or '',
            'subject': core_properties.subject or '',
            'keywords': core_properties.keywords or '',
            'comments': core_properties.comments or '',
            'created': core_properties.created.strftime('%Y-%m-%d %H:%M:%S') if core_properties.created else '',
            'modified': core_properties.modified.strftime('%Y-%m-%d %H:%M:%S') if core_properties.modified else '',
            'last_modified_by': core_properties.last_modified_by or '',
            'revision': core_properties.revision or 0,
            'category': core_properties.category or ''
        }
        
        metadata['word_count'] = self.get_word_count()
        metadata['paragraph_count'] = len(self.document.paragraphs)
        return metadata

    def get_paragraphs(self) -> List[str]:
        """
        获取DOCX文件的所有段落
        :return: 包含所有段落文本的列表
        """
        if not self.document:
            self._open_docx()
        
        return [paragraph.text for paragraph in self.document.paragraphs if paragraph.text.strip()]

    def get_word_count(self) -> int:
        """
        获取DOCX文件的总字数
        :return: 总字数
        """
        if not self.document:
            self._open_docx()
        
        word_count = 0
        for paragraph in self.document.paragraphs:
            word_count += len(paragraph.text.split())
        
        return word_count

    def extract_styles(self) -> List[str]:
        """
        提取文档中使用的样式信息
        :return: 包含所有唯一样式名称的列表
        """
        if not self.document:
            self._open_docx()
        
        styles = set()
        for paragraph in self.document.paragraphs:
            if paragraph.style.name:
                styles.add(paragraph.style.name)
        
        return list(styles)